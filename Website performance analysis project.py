from docx import Document

doc = Document()
doc.add_heading("Website Performance Analysis Report", 0)

text = """Website performance analysis including traffic, engagement,
frontend and backend architecture, and recommendations."""

for i in range(130):
    doc.add_heading(f"Section {i+1}", 1)
    doc.add_paragraph(text * 20)  # makes it long like real pages

doc.save("report.docx")